import io
import os
import tempfile
from pathlib import Path

import streamlit as st
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches


st.set_page_config(
    page_title="Proposal Cover Page Generator",
    page_icon="📄",
    layout="wide"
)


NAVY = "#0B2B68"
GREEN = "#63B52A"
ORANGE = "#F0A500"
LIGHT_GREEN = "#EAF4DE"


FONT_BOLD = [
    "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf",
    "/usr/share/fonts/truetype/liberation2/LiberationSerif-Bold.ttf"
]

FONT_REGULAR = [
    "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf",
    "/usr/share/fonts/truetype/liberation2/LiberationSerif-Regular.ttf"
]


def get_font(size, bold=True):
    paths = FONT_BOLD if bold else FONT_REGULAR

    for path in paths:
        if os.path.exists(path):
            return ImageFont.truetype(path, size=size)

    return ImageFont.load_default()


def fit_font(text, max_width, start_size=60, min_size=18, bold=True):
    text = str(text)

    size = start_size

    while size >= min_size:
        font = get_font(size, bold)

        box = font.getbbox(text)
        width = box[2] - box[0]

        if width <= max_width:
            return font

        size -= 2

    return get_font(min_size, bold)


def wrap_text(draw, text, font, max_width):
    text = str(text)

    words = text.split()

    lines = []
    current = ""

    for word in words:

        trial = word if not current else current + " " + word

        box = draw.textbbox(
            (0, 0),
            trial,
            font=font
        )

        width = box[2] - box[0]

        if width <= max_width:
            current = trial

        else:

            if current:
                lines.append(current)

            current = word

    if current:
        lines.append(current)

    return lines


def rounded_rectangle(
    draw,
    xy,
    radius,
    fill,
    outline=None,
    width=1
):
    draw.rounded_rectangle(
        xy,
        radius=radius,
        fill=fill,
        outline=outline,
        width=width
    )


def fit_crop(img, size):

    img = img.convert("RGB")

    target_w, target_h = size

    ratio = max(
        target_w / img.width,
        target_h / img.height
    )

    new_w = int(img.width * ratio)
    new_h = int(img.height * ratio)

    img = img.resize(
        (new_w, new_h),
        Image.Resampling.LANCZOS
    )

    left = (new_w - target_w) // 2
    top = (new_h - target_h) // 2

    return img.crop(
        (
            left,
            top,
            left + target_w,
            top + target_h
        )
    )


def fit_contain(img, size, bg=(255, 255, 255)):

    img = img.convert("RGBA")

    target_w, target_h = size

    ratio = min(
        target_w / img.width,
        target_h / img.height
    )

    new_w = max(
        1,
        int(img.width * ratio)
    )

    new_h = max(
        1,
        int(img.height * ratio)
    )

    img = img.resize(
        (new_w, new_h),
        Image.Resampling.LANCZOS
    )

    canvas = Image.new(
        "RGBA",
        size,
        bg + (255,)
    )

    canvas.alpha_composite(
        img,
        (
            (target_w - new_w) // 2,
            (target_h - new_h) // 2
        )
    )

    return canvas


def draw_image_collage(
    canvas,
    images,
    x,
    y,
    w,
    h
):

    draw = ImageDraw.Draw(canvas)

    if not images:

        rounded_rectangle(
            draw,
            (
                x,
                y,
                x + w,
                y + h
            ),
            35,
            "#F2F2F2",
            NAVY,
            6
        )

        text = "UPLOAD PROJECT IMAGES"

        font = get_font(34)

        box = draw.textbbox(
            (0, 0),
            text,
            font=font
        )

        text_w = box[2] - box[0]
        text_h = box[3] - box[1]

        draw.text(
            (
                x + (w - text_w) / 2,
                y + (h - text_h) / 2
            ),
            text,
            font=font,
            fill=NAVY
        )

        return

    n = len(images)

    gap = 12

    if n == 1:

        tiles = [
            (
                x,
                y,
                w,
                h
            )
        ]

    elif n == 2:

        left_w = (w - gap) // 2

        tiles = [
            (
                x,
                y,
                left_w,
                h
            ),
            (
                x + left_w + gap,
                y,
                w - left_w - gap,
                h
            )
        ]

    elif n == 3:

        left_w = int(w * 0.47)

        right_x = x + left_w + gap

        right_w = w - left_w - gap

        half_h = (h - gap) // 2

        tiles = [
            (
                x,
                y,
                left_w,
                h
            ),
            (
                right_x,
                y,
                right_w,
                half_h
            ),
            (
                right_x,
                y + half_h + gap,
                right_w,
                half_h
            )
        ]

    else:

        left_w = int(w * 0.46)

        right_x = x + left_w + gap

        right_w = w - left_w - gap

        half_h = (h - gap) // 2

        tiles = [
            (
                x,
                y,
                left_w,
                h
            ),
            (
                right_x,
                y,
                right_w,
                half_h
            ),
            (
                right_x,
                y + half_h + gap,
                right_w,
                half_h
            )
        ]

    for i, box in enumerate(tiles):

        image = images[i % len(images)]

        tile_w = box[2]
        tile_h = box[3]

        tile = fit_crop(
            image,
            (
                tile_w,
                tile_h
            )
        )

        mask = Image.new(
            "L",
            (
                tile_w,
                tile_h
            ),
            0
        )

        mask_draw = ImageDraw.Draw(mask)

        mask_draw.rounded_rectangle(
            (
                0,
                0,
                tile_w - 1,
                tile_h - 1
            ),
            radius=28,
            fill=255
        )

        canvas.paste(
            tile,
            (
                box[0],
                box[1]
            ),
            mask
        )


def build_cover(
    project_type,
    project_name,
    reference,
    date_text,
    client_name,
    client_logo,
    funding_name,
    funding_logo,
    project_images,
    companies
):

    W = 2048
    H = 2816

    canvas = Image.new(
        "RGB",
        (
            W,
            H
        ),
        "white"
    )

    draw = ImageDraw.Draw(canvas)

    for yy in range(H):

        t = yy / H

        c = int(
            255 - 14 * t
        )

        draw.line(
            (
                0,
                yy,
                W,
                yy
            ),
            fill=(
                c,
                min(255, c + 1),
                255
            )
        )

    draw.polygon(
        [
            (1550, 0),
            (W, 0),
            (W, 900),
            (1760, 690)
        ],
        fill="#EDF4FA"
    )

    draw.polygon(
        [
            (1780, 0),
            (W, 0),
            (W, 760)
        ],
        fill="#E7F0F8"
    )

    draw.polygon(
        [
            (0, 1180),
            (350, 880),
            (650, 1150),
            (300, 1500)
        ],
        fill="#F4F8FB"
    )

    if client_logo:

        logo = fit_contain(
            client_logo,
            (
                230,
                180
            )
        )

        canvas.paste(
            logo,
            (
                100,
                120
            ),
            logo
        )

    if funding_logo:

        logo = fit_contain(
            funding_logo,
            (
                300,
                190
            )
        )

        canvas.paste(
            logo,
            (
                W - 400,
                105
            ),
            logo
        )

    heading = str(
        project_type
    ).upper()

    hf = fit_font(
        heading,
        880,
        76,
        44,
        True
    )

    hb = draw.textbbox(
        (0, 0),
        heading,
        font=hf
    )

    hw = hb[2] - hb[0]

    hx = (W - hw) // 2

    hy = 180

    draw.polygon(
        [
            (
                hx - 80,
                hy + 20
            ),
            (
                hx + hw + 80,
                hy + 20
            ),
            (
                hx + hw + 105,
                hy + 130
            ),
            (
                hx - 55,
                hy + 130
            )
        ],
        fill="#F6D78E"
    )

    draw.polygon(
        [
            (
                hx - 70,
                hy
            ),
            (
                hx + hw + 70,
                hy
            ),
            (
                hx + hw + 95,
                hy + 100
            ),
            (
                hx - 45,
                hy + 100
            )
        ],
        fill="white",
        outline=ORANGE
    )

    draw.text(
        (
            hx,
            hy + 13
        ),
        heading,
        font=hf,
        fill=NAVY
    )

    bx = 30
    by = 490
    bw = W - 60
    bh = 430

    rounded_rectangle(
        draw,
        (
            bx,
            by,
            bx + bw,
            by + bh
        ),
        58,
        LIGHT_GREEN,
        NAVY,
        8
    )

    project_name = str(
        project_name
    ).strip()

    if not project_name:
        project_name = "PROJECT NAME"

    pf = fit_font(
        project_name,
        bw - 90,
        54,
        28,
        True
    )

    lines = wrap_text(
        draw,
        project_name,
        pf,
        bw - 90
    )

    box = pf.getbbox("Ag")

    line_h = (
        box[3] -
        box[1] +
        8
    )

    ty = by + 55

    for line in lines[:6]:

        draw.text(
            (
                bx + 45,
                ty
            ),
            line,
            font=pf,
            fill=NAVY
        )

        ty += line_h

    draw.line(
        (
            bx + 50,
            by + bh - 118,
            bx + bw - 50,
            by + bh - 118
        ),
        fill="#222222",
        width=4
    )

    reference_text = (
        f"Reference No: {reference}"
    )

    rf = fit_font(
        reference_text,
        1200,
        46,
        28,
        True
    )

    rb = draw.textbbox(
        (0, 0),
        reference_text,
        font=rf
    )

    rw = rb[2] - rb[0]

    draw.text(
        (
            (W - rw) // 2,
            by + bh - 105
        ),
        reference_text,
        font=rf,
        fill="#41622D"
    )

    date_text = str(
        date_text
    ).strip()

    if not date_text:
        date_text = "Aug, 2026"

    df = fit_font(
        date_text,
        360,
        42,
        28,
        True
    )

    db = draw.textbbox(
        (0, 0),
        date_text,
        font=df
    )

    dw = db[2] - db[0]

    rounded_rectangle(
        draw,
        (
            (W - dw) // 2 - 55,
            870,
            (W + dw) // 2 + 55,
            965
        ),
        15,
        NAVY,
        "white",
        3
    )

    draw.text(
        (
            (W - dw) // 2,
            885
        ),
        date_text,
        font=df,
        fill="white"
    )

    draw.rectangle(
        (
            0,
            1190,
            W,
            1225
        ),
        fill="#2C7EC7"
    )

    draw_image_collage(
        canvas,
        project_images,
        35,
        1235,
        W - 70,
        930
    )

    sy = 2235

    rounded_rectangle(
        draw,
        (
            15,
            sy,
            W - 15,
            H - 25
        ),
        38,
        "white",
        NAVY,
        6
    )

    draw.polygon(
        [
            (
                W // 2 - 270,
                sy + 55
            ),
            (
                W // 2 + 270,
                sy + 55
            ),
            (
                W // 2 + 220,
                sy + 115
            ),
            (
                W // 2 - 220,
                sy + 115
            )
        ],
        fill=NAVY
    )

    submitted_text = "SUBMITTED BY"

    sf = get_font(
        46,
        True
    )

    sb = draw.textbbox(
        (0, 0),
        submitted_text,
        font=sf
    )

    submitted_w = (
        sb[2] -
        sb[0]
    )

    draw.text(
        (
            (W - submitted_w) // 2,
            sy + 64
        ),
        submitted_text,
        font=sf,
        fill="white"
    )

    if not companies:
        companies = [
            {
                "name": "",
                "role": "LEAD",
                "logo": None
            }
        ]

    companies = companies[:5]

    n = len(companies)

    col_w = (
        W - 90
    ) // n

    for i, company in enumerate(companies):

        cx = 45 + i * col_w

        if i:

            draw.line(
                (
                    cx,
                    sy + 155,
                    cx,
                    H - 70
                ),
                fill="#A9A9A9",
                width=2
            )

        logo = company.get(
            "logo"
        )

        if logo:

            logo_image = fit_contain(
                logo,
                (
                    col_w - 60,
                    190
                )
            )

            canvas.paste(
                logo_image,
                (
                    cx + 30,
                    sy + 165
                ),
                logo_image
            )

        name = str(
            company.get(
                "name",
                ""
            )
        ).strip()

        if not name:
            name = "COMPANY NAME"

        role = str(
            company.get(
                "role",
                "LEAD"
            )
        )

        nf = fit_font(
            name,
            col_w - 55,
            32,
            18,
            True
        )

        name_lines = wrap_text(
            draw,
            name,
            nf,
            col_w - 55
        )

        ny = sy + 365

        for line in name_lines[:2]:

            nb = draw.textbbox(
                (0, 0),
                line,
                font=nf
            )

            name_w = (
                nb[2] -
                nb[0]
            )

            draw.text(
                (
                    cx +
                    (col_w - name_w) // 2,
                    ny
                ),
                line,
                font=nf,
                fill=NAVY
            )

            ny += 38

        role_text = (
            f"[ {role.upper()} ]"
        )

        role_font = get_font(
            30,
            True
        )

        rb = draw.textbbox(
            (0, 0),
            role_text,
            font=role_font
        )

        role_w = (
            rb[2] -
            rb[0]
        )

        draw.text(
            (
                cx +
                (col_w - role_w) // 2,
                ny + 8
            ),
            role_text,
            font=role_font,
            fill=GREEN
        )

    return canvas


def image_from_upload(upload):

    if upload is None:
        return None

    try:
        return Image.open(
            upload
        ).convert("RGBA")
    except Exception:
        return None


def make_docx(image):

    output = io.BytesIO()

    with tempfile.NamedTemporaryFile(
        suffix=".png",
        delete=False
    ) as tmp:

        image.save(
            tmp.name,
            format="PNG",
            dpi=(300, 300)
        )

        tmp_path = tmp.name

    try:

        doc = Document()

        section = doc.sections[0]

        section.page_width = Inches(8.5)
        section.page_height = Inches(11.69)

        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)

        paragraph = doc.paragraphs[0]

        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0

        run = paragraph.add_run()

        run.add_picture(
            tmp_path,
            width=Inches(8.5)
        )

        doc.save(output)

    finally:

        try:
            os.unlink(tmp_path)
        except Exception:
            pass

    output.seek(0)

    return output


st.title(
    "📄 Proposal Cover Page Generator"
)

st.caption(
    "Create professional proposal and EOI cover pages using a reusable template."
)


with st.sidebar:

    st.header(
        "1. Project Details"
    )

    project_type = st.selectbox(
        "Document Type",
        [
            "Expression of Interest",
            "Request for Proposal",
            "Technical Proposal",
            "Financial Proposal",
            "Proposal"
        ]
    )

    project_name = st.text_area(
        "Project Name",
        height=140,
        value=(
            "Consulting Services for Maritime Diagnosis and "
            "Modelling, Study and Alternatives, Executive Project "
            "and Environmental and Social Studies, for the "
            "Construction of the Coastal Protection Infrastructure, "
            "Rehabilitation and Requalification of the Marginal "
            "of the Municipality of the City of Vilankulo, "
            "Inhambane Province"
        )
    )

    reference = st.text_input(
        "Reference No.",
        "MZ-MEF-DNT-556716-CS-QCBS"
    )

    date_text = st.text_input(
        "Date",
        "Aug, 2026"
    )

    client_name = st.text_input(
        "Client / Government",
        ""
    )

    funding_name = st.text_input(
        "Funding / Bank",
        ""
    )

    st.header(
        "2. Logos"
    )

    client_logo_file = st.file_uploader(
        "Client / Government Logo",
        type=[
            "png",
            "jpg",
            "jpeg"
        ],
        key="client_logo"
    )

    funding_logo_file = st.file_uploader(
        "Funding / Bank Logo",
        type=[
            "png",
            "jpg",
            "jpeg"
        ],
        key="funding_logo"
    )

    st.header(
        "3. Project Images"
    )

    image_files = st.file_uploader(
        "Upload 1–4 Project Images",
        type=[
            "png",
            "jpg",
            "jpeg"
        ],
        accept_multiple_files=True,
        key="project_images"
    )

    if image_files:
        image_files = image_files[:4]

    st.header(
        "4. Submitted By"
    )

    company_count = st.number_input(
        "Number of Organizations",
        min_value=1,
        max_value=5,
        value=3,
        step=1
    )

    roles = [
        "LEAD",
        "JV",
        "SUB-CON",
        "ASSOCIATE",
        "PARTNER"
    ]

    companies = []

    for i in range(
        int(company_count)
    ):

        st.markdown(
            f"### Organization {i + 1}"
        )

        company_name = st.text_input(
            "Company Name",
            key=f"company_name_{i}"
        )

        company_role = st.selectbox(
            "Role",
            roles,
            key=f"company_role_{i}"
        )

        company_logo = st.file_uploader(
            "Company Logo",
            type=[
                "png",
                "jpg",
                "jpeg"
            ],
            key=f"company_logo_{i}"
        )

        companies.append(
            {
                "name": company_name,
                "role": company_role,
                "logo": image_from_upload(
                    company_logo
                )
            }
        )

    generate = st.button(
        "🚀 Generate Cover",
        type="primary",
        use_container_width=True
    )


client_logo = image_from_upload(
    client_logo_file
)

funding_logo = image_from_upload(
    funding_logo_file
)


project_images = []

if image_files:

    for uploaded_file in image_files:

        image = image_from_upload(
            uploaded_file
        )

        if image:
            project_images.append(
                image
            )


cover = build_cover(
    project_type=project_type,
    project_name=project_name,
    reference=reference,
    date_text=date_text,
    client_name=client_name,
    client_logo=client_logo,
    funding_name=funding_name,
    funding_logo=funding_logo,
    project_images=project_images,
    companies=companies
)


left_column, right_column = st.columns(
    [
        1.15,
        0.85
    ]
)


with left_column:

    st.subheader(
        "Live Preview"
    )

    st.image(
        cover,
        use_container_width=True
    )


with right_column:

    st.subheader(
        "Output"
    )

    docx_file = make_docx(
        cover
    )

    st.download_button(
        "⬇️ Download Word Cover Page",
        data=docx_file.getvalue(),
        file_name="Proposal_Cover_Page.docx",
        mime=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        use_container_width=True
    )

    png_file = io.BytesIO()

    cover.save(
        png_file,
        format="PNG",
        dpi=(300, 300)
    )

    st.download_button(
        "🖼️ Download PNG Cover",
        data=png_file.getvalue(),
        file_name="Proposal_Cover_Page.png",
        mime="image/png",
        use_container_width=True
    )

    st.markdown("---")

    st.info(
        "The current version creates the complete cover as a "
        "high-resolution image and places it into Word. This "
        "keeps the layout consistent across different computers."
    )
